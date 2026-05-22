from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
WORD = ROOT / "docs" / "word"
DATE = "2026 年 5 月 21 日"

CRITERIA = [
    ("C01", "行情延迟提示", "行情可能延迟或接口返回缓存数据，界面必须显示最近刷新时间和数据来源。", "顶部刷新状态显示具体时间；行情异常时提示“显示上次数据”。"),
    ("C02", "涨跌颜色一致", "股票软件常用红涨绿跌，但不同市场习惯可能不同，颜色含义必须稳定。", "全局统一红色表示上涨、绿色表示下跌，并在原型说明中写明。"),
    ("C03", "价格单位清晰", "成交量、成交额、价格单位不清会导致商科用户误读。", "成交量显示“手/万手”，成交额显示“万/亿”，价格显示“元”。"),
    ("C04", "股票代码防错", "用户可能输入不存在、重复或格式错误的股票代码。", "添加自选股时校验非空、重复；后续实现增加 A 股代码格式校验。"),
    ("C05", "自选股为空状态", "首次使用或删除全部股票后，界面不能空白无解释。", "显示空状态提示，引导用户添加第一只股票。"),
    ("C06", "删除防误触", "用户可能误删自选股和相关提醒规则。", "删除前增加确认提示，删除后可在短时间内撤销。"),
    ("C07", "提醒规则可理解", "涨超、跌超、涨过、跌破容易混淆。", "提醒类型使用清晰中文标签，并在阈值输入旁显示单位。"),
    ("C08", "提醒避免骚扰", "价格持续满足条件时可能重复弹窗。", "提醒触发后标记状态，同一规则不连续重复弹窗。"),
    ("C09", "阈值输入防错", "用户可能输入负数、空值或非数字。", "保存提醒时进行数值校验，非法输入给出明确提示。"),
    ("C10", "图表可读性", "折线图点过密或颜色太浅会影响趋势判断。", "限制展示点数，使用高对比折线，并标注最高/最低。"),
    ("C11", "分时/日 K 状态明确", "用户可能不知道当前看的是分时还是日 K。", "标签选中态明显，图表标题同步变化。"),
    ("C12", "行情刷新不打断操作", "自动刷新时不应清空用户正在输入的提醒规则。", "刷新只更新行情区域，不重置表单输入。"),
    ("C13", "网络失败降级", "公开行情接口可能不可用，不能导致程序崩溃。", "保留模拟数据和本地缓存，失败时显示错误说明。"),
    ("C14", "主界面信息不过载", "股票盯盘容易堆积太多指标，反而降低效率。", "首页只放现价、涨跌、开高低、成交量额等核心字段。"),
    ("C15", "资产/持仓语义", "商科用户希望能把自选股理解为资产观察清单。", "自选股区域命名为观察清单，后续可扩展持仓备注。"),
    ("C16", "多屏适配", "用户可能在笔记本、外接显示器或投影环境演示。", "原型使用响应式布局，窄屏时右侧提醒区下移。"),
    ("C17", "中文表达专业但不晦涩", "金融术语过多会增加学习成本。", "保留常用术语，避免 Level-2、复权等本迭代不需要概念。"),
    ("C18", "关键操作反馈", "添加、删除、刷新、保存提醒后用户需要知道是否成功。", "通过列表变化、刷新时间、提醒列表新增和高亮状态反馈。"),
    ("C19", "演示可控性", "课堂演示时真实行情波动不可控或无网络。", "内置模拟行情，确保离线也能展示完整流程。"),
    ("C20", "需求追踪", "评审者需要看到界面元素与需求/用例的对应关系。", "界面原型文档维护需求-用例-界面映射表。"),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if fill:
        shade(cell, fill)
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, "EAF2F8")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()


def add_history_row(doc, description):
    if not doc.tables:
        return
    table = doc.tables[0]
    row = table.add_row().cells
    values = [DATE, "1.2", description, "项目组"]
    for i, value in enumerate(values):
        set_cell(row[i], value)


def save(doc, path):
    doc.save(path)
    print(path)


def update_ui_prototype():
    path = WORD / "StockMonitor-界面原型.docx"
    doc = Document(path)
    add_history_row(doc, "根据软件界面 checklist 补充 20 条股票盯盘用户问题与界面准则")

    doc.add_heading("7. 股票盯盘用户可能遇到的问题与 20 条界面准则", level=1)
    doc.add_paragraph("依据课程给出的软件界面 checklist，结合股票盯盘场景补充以下 20 条准则。每条准则都对应一个真实用户可能遇到的问题，并给出本项目原型或后续实现中的改进方向。")
    add_table(doc, ["编号", "准则", "用户可能遇到的问题", "本项目改进方式"], CRITERIA)

    doc.add_heading("8. V1.2 原型修订结论", level=1)
    for item in [
        "将刷新时间、数据来源、模拟行情兜底作为行情类界面的必要反馈。",
        "将输入校验、删除确认、提醒防重复作为后续技术实现的高优先级改进。",
        "保持主界面简洁，不引入真实交易、复杂蜡烛图和过多金融术语。",
        "继续维护需求-用例-界面映射，便于课程评审追踪。"
    ]:
        doc.add_paragraph(item, style="List Bullet")
    save(doc, path)


def update_iteration_plan():
    path = WORD / "StockMonitor-hw1迭代计划-v1.1.docx"
    doc = Document(path)
    add_history_row(doc, "新增 V1.2 界面准则迭代：补充 20 条股票盯盘用户问题与改进任务")

    doc.add_heading("8. V1.2 界面准则增强迭代", level=1)
    doc.add_paragraph("在完成 V1.1 界面原型后，项目组根据软件界面 checklist 继续进行小步迭代。本次迭代不增加大功能，而是从股票盯盘用户角度补充风险点和可用性准则，使原型更贴近真实使用。")
    add_table(doc, ["任务", "说明", "优先级", "输出"], [
        ["补充用户问题清单", "围绕行情延迟、提醒误触、输入错误、图表误读等场景扩展 20 条准则。", "高", "界面原型文档第 7 节"],
        ["更新修订历史", "所有被完善的文档补充 1.2 版本记录，体现持续迭代。", "高", "修订历史记录"],
        ["完善需求风险应对", "把数据源失败、提醒骚扰、删除误触等问题纳入风险管理。", "高", "迭代评估报告"],
        ["保持范围克制", "不新增真实交易功能，只提升盯盘体验和原型可解释性。", "中", "迭代计划结论"],
    ])
    save(doc, path)


def update_evaluation():
    path = WORD / "StockMonitor-hw1迭代评估报告-v1.1.docx"
    doc = Document(path)
    add_history_row(doc, "补充 V1.2 小组评审记录：按 20 条界面准则继续完善原型")

    doc.add_heading("7. V1.2 追加评审：20 条界面准则处理结果", level=1)
    add_table(doc, ["类别", "发现的问题", "修订或后续处理"], [
        ["行情反馈", "用户可能不知道行情是否最新、是否来自真实接口。", "界面准则增加刷新时间、数据来源和缓存提示。"],
        ["输入防错", "股票代码、提醒阈值可能输错。", "将代码格式、重复添加、阈值数值校验列为后续 Must。"],
        ["提醒体验", "同一规则可能频繁弹窗打扰用户。", "提醒触发后标记状态，避免连续重复提示。"],
        ["图表理解", "用户可能分不清分时和日 K，或误读单位。", "要求标签选中态明显，并补充单位和最高/最低标注。"],
        ["演示稳定", "课堂演示可能无网络或行情接口不可用。", "保留模拟行情作为演示兜底，不牺牲验收质量。"],
    ])

    doc.add_heading("8. V1.2 经验总结", level=1)
    for item in [
        "界面 checklist 不只是美观检查，也能发现业务场景中的真实风险。",
        "股票盯盘软件的关键不是堆指标，而是让用户知道数据是否可信、操作是否成功、提醒是否可靠。",
        "持续小步迭代比一次性写完文档更符合课程对迭代开发过程的要求。"
    ]:
        doc.add_paragraph(item, style="List Bullet")
    save(doc, path)


def update_srs_addendum():
    path = WORD / "StockMonitor-SRS补充说明.docx"
    doc = Document(path)
    add_history_row(doc, "补充 20 条界面准则对应的需求风险说明")

    doc.add_heading("5. V1.2 新增界面准则对应的需求风险", level=1)
    add_table(doc, ["需求风险", "对应准则编号", "SRS 影响"], [
        ["行情数据可信度不足", "C01、C13、C19", "补充刷新时间、数据来源、缓存和模拟数据兜底要求。"],
        ["用户输入错误导致规则无效", "C04、C07、C09", "补充股票代码和提醒阈值校验需求。"],
        ["提醒过度打扰用户", "C08、C18", "补充提醒触发状态和反馈要求。"],
        ["趋势图被误读", "C03、C10、C11", "补充单位、标签状态和图表可读性要求。"],
        ["评审难以追踪需求", "C20", "继续维护需求-用例-界面映射。"],
    ])
    save(doc, path)


if __name__ == "__main__":
    update_ui_prototype()
    update_iteration_plan()
    update_evaluation()
    update_srs_addendum()
