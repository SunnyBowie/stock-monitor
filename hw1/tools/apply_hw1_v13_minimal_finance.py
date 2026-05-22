from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
WORD = ROOT / "docs" / "word"
DATE = "2026 年 5 月 22 日"
VERSION = "1.3"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    for run_item in paragraph.runs:
        run_item.font.name = "Microsoft YaHei"
        run_item._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run_item.font.size = Pt(9)
    if fill:
        shade(cell, fill)
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell(t.rows[0].cells[index], header, True, "EAF2F8")
    for row in rows:
        cells = t.add_row().cells
        for index, value in enumerate(row):
            set_cell(cells[index], value)
    doc.add_paragraph()
    return t


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_history_row(doc, description):
    if not doc.tables:
        return
    row = doc.tables[0].add_row().cells
    for index, value in enumerate([DATE, VERSION, description, "项目组"]):
        set_cell(row[index], value)


def save(doc, path):
    doc.save(path)
    print(path)


def update_vision():
    path = WORD / "StockMonitor-前景文档.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：调整为简约看盘定位，补充财务报表可视化愿景")

    doc.add_heading("6. V1.3 简约看盘产品定位补充", level=1)
    doc.add_paragraph(
        "本轮迭代将 Stock Monitor 从“轻量股票盯盘工具”进一步收敛为“界面简约的股票看盘软件”。"
        "简约不等于功能少，而是将信息层级做减法：第一屏解决日常盯盘，第二层承载趋势、盘口和财务深度。"
    )
    table(doc, ["设计原则", "产品含义", "界面落点"], [
        ["1 屏解决日常", "用户打开软件后无需在大量指标中寻找重点。", "自选股首页只显示股票简称、现价、涨跌幅。"],
        ["2 层解决深度", "用户主动点进个股后才展示更复杂的信息。", "详情页展示分时/K 线、关键盘口、财务报表和新闻。"],
        ["不打扰交互", "刷新和提醒不抢夺用户注意力。", "价格变动轻微闪烁，提醒触发后标记状态避免重复打扰。"],
        ["财报反表格化", "财务数据用于辅助判断，不复刻密集报表。", "近四季度营收/净利润柱状图、PE 分位进度条、ROE 等摘要字段。"],
    ])

    doc.add_heading("6.1 财务报表能力边界", level=2)
    bullets(doc, [
        "第一阶段只展示财务摘要，不提供完整资产负债表、利润表和现金流量表明细。",
        "财务模块采用可视化表达：营收和净利润用柱状图，估值水位用进度条，核心指标用小卡片。",
        "财务数据刷新频率低于行情数据，可按日或按财报期缓存，避免影响实时盯盘性能。",
        "财务信息只做观察辅助，不生成投资建议或买卖信号。"
    ])
    save(doc, path)


def update_srs():
    path = WORD / "StockMonitor-软件需求规约SRS.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：补充简约信息层级、K 线和财务报表摘要需求")

    doc.add_heading("4. V1.3 简约看盘需求补充", level=1)
    doc.add_paragraph(
        "V1.3 需求补充强调信息层级：系统首页必须保持低噪声，只展示完成日常盯盘所需的最小字段；"
        "个股详情页作为第二层，按用户主动选择展示图表、盘口和财务摘要。"
    )
    table(doc, ["编号", "需求", "优先级", "验收标准"], [
        ["FR-09", "首页自选股列表只展示股票简称、现价、涨跌幅。", "高", "列表中不出现成交额、盘口、新闻等次级字段。"],
        ["FR-10", "系统提供分时图和 K 线图切换。", "高", "图表标签清晰，分时显示时间轴，K 线显示交易日轴。"],
        ["FR-11", "系统展示简化财务报表摘要。", "中", "至少展示近 4 季度营收、净利润、ROE、负债/资本或现金流摘要。"],
        ["FR-12", "财务报表采用图形化方式呈现。", "中", "营收/净利润用柱状图，估值水位用进度条，不以密集表格作为默认展示。"],
        ["FR-13", "系统提供极简模式。", "中", "涨跌只通过文字颜色表达，不使用大面积红绿底色。"],
    ])

    doc.add_heading("4.1 信息层级约束", level=2)
    table(doc, ["层级", "展示内容", "不展示内容"], [
        ["首页/自选", "股票简称、现价、涨跌幅、刷新状态。", "盘口明细、财报明细、新闻列表、技术指标窗口。"],
        ["详情上半区", "大字体价格、涨跌额/涨跌幅、分时图、K 线切换。", "多窗口技术指标平铺。"],
        ["详情下半区", "关键数据、买一/卖一、财务摘要、简明新闻。", "完整财务报表、十档盘口、复杂排行榜。"],
        ["隐藏抽屉", "MA、MACD、RSI、KDJ 等可选指标。", "默认同时展开 3 个以上指标窗口。"],
    ])
    save(doc, path)


def update_use_cases():
    path = WORD / "StockMonitor-用例规约.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：补充查看财务报表摘要用例")

    doc.add_heading("UC-07 查看财务报表摘要", level=1)
    doc.add_heading("简要说明", level=2)
    doc.add_paragraph(
        "用户在个股详情页查看简化财务报表摘要，用于快速理解公司的盈利能力、现金流和估值水位。"
        "该用例强调图形化和低密度展示，不要求用户阅读完整财务报表。"
    )
    doc.add_heading("基本流", level=2)
    for item in [
        "用户在自选股列表中选择一只股票。",
        "系统进入个股详情页，并加载该股票的财务摘要缓存。",
        "系统展示近四季度营收和净利润柱状图。",
        "系统展示 PE 历史分位进度条。",
        "系统展示 ROE、负债/资本、经营现金流等核心摘要字段。"
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("备选流", level=2)
    bullets(doc, [
        "财务数据源不可用：系统显示最近一次缓存，并标记数据日期。",
        "部分字段缺失：系统隐藏缺失字段或显示“暂无”，不得影响行情和图表查看。",
        "用户不需要财务信息：财务模块保持在第二层或折叠区域，不干扰首页盯盘。"
    ])
    doc.add_heading("特殊需求", level=2)
    doc.add_paragraph("财务摘要刷新不得阻塞行情刷新；财务数据不得被解释为投资建议。")
    doc.add_heading("前置条件", level=2)
    doc.add_paragraph("用户已选中一只股票。")
    doc.add_heading("后置条件", level=2)
    doc.add_paragraph("详情页展示可用财务摘要或明确的暂无/缓存状态。")
    save(doc, path)


def update_ui_doc():
    path = WORD / "StockMonitor-界面原型.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：更新为简约风界面，新增财务报表模块说明")

    doc.add_heading("9. V1.3 简约风界面原型说明", level=1)
    doc.add_paragraph(
        "V1.3 原型采用暗黑简约风格，核心不是减少功能，而是减少默认暴露的信息。"
        "用户日常在自选股首页完成观察，需要深度信息时再进入个股详情。"
    )
    table(doc, ["界面区域", "保留信息", "交互方式"], [
        ["自选股列表", "股票简称、代码、现价、涨跌幅。", "点击进入详情；左/右滑动触发删除或置顶。"],
        ["个股详情上半区", "大字体现价、涨跌额、涨跌幅、刷新时间、分时/K 线。", "分段按钮切换图表周期。"],
        ["关键数据", "今开、昨收、最高、最低、成交量、成交额、换手。", "只在详情页展示，不进入首页列表。"],
        ["盘口", "买一、卖一和成交量。", "只展示最小盘口，不展示十档。"],
        ["财务报表", "营收/净利润柱状图、PE 分位、ROE、负债/资本、经营现金流。", "图形化展示，避免密集表格。"],
        ["发现页", "搜索框、上证/深证/创业板指数卡片。", "不做复杂排行榜。"],
    ])

    doc.add_heading("9.1 财务报表界面设计", level=2)
    bullets(doc, [
        "财务报表放在个股详情的第二层，默认不出现在自选股首页。",
        "营收和净利润使用近四季度柱状图表达趋势。",
        "PE 历史分位使用进度条表达估值水位。",
        "毛利率、ROE、负债/资本、经营现金流使用小卡片表达，避免用户阅读完整报表。"
    ])
    save(doc, path)


def update_iteration_plan():
    path = WORD / "StockMonitor-hw1迭代计划-v1.1.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：简约风与财务报表原型迭代计划")

    doc.add_heading("9. V1.3 简约风与财务报表迭代", level=1)
    table(doc, ["任务", "说明", "优先级", "输出"], [
        ["重构信息层级", "首页只保留简称、现价、涨跌幅，详情页承载深度信息。", "高", "prototype/index.html"],
        ["补充分时/K 线时间轴", "分时图标注 09:30、11:30、13:00、15:00，K 线标注交易日。", "高", "界面原型文档、SRS"],
        ["新增财务报表模块", "用柱状图和进度条展示财务摘要。", "中", "界面原型、SRS、用例规约"],
        ["更新用例模型", "新增“查看财务报表摘要”用例。", "中", "use_case.puml、用例规约"],
        ["同步数据源分析", "补充财务摘要字段、刷新频率和缓存策略。", "中", "data-source-analysis.md"],
    ])
    save(doc, path)


def update_evaluation():
    path = WORD / "StockMonitor-hw1迭代评估报告-v1.1.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：评估简约风原型与财务报表模块")

    doc.add_heading("9. V1.3 评估：简约风与财务报表", level=1)
    table(doc, ["评估项", "结果", "说明"], [
        ["首页信息密度", "通过", "自选股首页只保留股票简称、现价、涨跌幅，符合简约定位。"],
        ["深度信息承载", "通过", "个股详情页承载分时/K 线、盘口和财务摘要，符合 2 层解决深度原则。"],
        ["财务报表表达", "通过", "使用柱状图、进度条和摘要卡片，避免传统财务表格的高密度数字。"],
        ["后续实现风险", "可控", "财务摘要为低频数据，可通过缓存和模拟数据保证演示稳定。"],
    ])
    doc.add_heading("9.1 后续风险", level=2)
    bullets(doc, [
        "真实财务数据字段在不同数据源中命名不同，hw2 需要通过 FinancialDataProvider 适配。",
        "PE 历史分位需要历史估值数据或本地缓存，课程版本可先使用模拟分位。",
        "财务信息不能被界面文案表达成投资建议，应始终定位为观察辅助。"
    ])
    save(doc, path)


def update_srs_addendum():
    path = WORD / "StockMonitor-SRS补充说明.docx"
    doc = Document(path)
    add_history_row(doc, "V1.3：补充简约风信息层级与财务报表需求追踪")

    doc.add_heading("6. V1.3 简约风与财务报表需求追踪", level=1)
    table(doc, ["新增需求", "原型验证方式", "追踪结论"], [
        ["首页信息精简", "自选列表只展示简称、现价、涨跌幅。", "已在 prototype/index.html 中实现。"],
        ["分时/K 线时间轴", "图表区在分时模式标注 09:30、11:30、13:00、15:00，K 线标注交易日。", "已在原型中体现。"],
        ["财务报表摘要", "详情页财务报表卡片展示营收/净利润、PE 分位和核心指标。", "已在原型中体现，后续由 FinancialDataProvider 支撑。"],
        ["发现页克制", "发现页只保留搜索和 3 个指数卡片。", "已在原型中体现。"],
    ])
    save(doc, path)


def main():
    update_vision()
    update_srs()
    update_use_cases()
    update_ui_doc()
    update_iteration_plan()
    update_evaluation()
    update_srs_addendum()


if __name__ == "__main__":
    main()
