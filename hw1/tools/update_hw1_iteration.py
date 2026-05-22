from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "word"
OUT.mkdir(parents=True, exist_ok=True)

PROJECT = "Stock Monitor Windows 端股票盯盘软件"
DATE = "2026 年 5 月 21 日"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if fill:
        shade(cell, fill)
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup(doc, subtitle):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(35, 100, 170)
    doc.styles["Heading 3"].font.size = Pt(11.5)
    doc.styles["Heading 3"].font.color.rgb = RGBColor(15, 139, 141)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(23, 54, 93)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.bold = True
    r.font.size = Pt(15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"版本 1.1    {DATE}")
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell(t.rows[0].cells[i], h, True, "EAF2F8")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell(cells[i], value)
    doc.add_paragraph()
    return t


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def history(doc, note):
    doc.add_heading("修订历史记录", level=1)
    table(doc, ["日期", "版本", "说明", "作者"], [[DATE, "1.1", note, "项目组"]])


def build_iteration_plan():
    doc = Document()
    setup(doc, "hw1 界面原型迭代计划")
    history(doc, "根据课程指导截图更新为界面原型迭代，采用滚动式规划并应对需求风险")

    doc.add_heading("1. 本迭代定位", level=1)
    doc.add_paragraph("hw1 是界面原型迭代，目标是应对需求风险：通过用户调研、Vision、SRS、用例模型和可交互界面原型，尽早验证核心需求是否有价值、界面是否可理解、功能范围是否可控。")
    bullets(doc, [
        "过程：采用迭代开发，覆盖需求、设计、原型实现、评审和改进。",
        "设计：模块化界面设计，主要界面元素可追溯到需求和用例。",
        "编码：原型采用 HTML/CSS/JavaScript，核心实现遵循清晰命名和模块化组织。",
        "测试：本迭代重点为人工功能测试和界面 checklist 检查。",
        "版本管理：使用 Git/GitHub 管理文档、模型和原型代码。"
    ])

    doc.add_heading("2. 小组分工与进度安排", level=1)
    table(doc, ["角色", "职责", "本迭代任务"], [
        ["组长/项目经理", "范围控制、进度协调、GitHub 管理", "制定迭代计划，组织评审和提交。"],
        ["需求分析", "用户调研、需求定义、SRS 和用例", "更新 Vision、SRS、用例图和核心用例规约。"],
        ["原型/前端", "界面设计与可交互原型", "实现 HTML 原型，说明界面元素与需求/用例关系。"],
        ["测试/文档", "界面检查、评估报告", "按 checklist 测试原型，整理评审记录和迭代评估。"],
    ])

    doc.add_heading("3. 5 版滚动式迭代方案", level=1)
    table(doc, ["版本", "目标", "需求风险应对", "交付物", "通过标准"], [
        ["V1 用户访谈版", "确认用户、问题和功能边界", "避免做成专业交易终端，控制功能数量", "用户画像、问题说明、Vision 草稿", "能回答谁用、解决什么问题、为什么有价值"],
        ["V2 需求建模版", "把高优先级需求转成用例和 SRS", "防止需求含糊导致后续原型返工", "SRS、PlantUML 用例图、核心用例规约", "主要需求均有用例支撑"],
        ["V3 低保真界面版", "确定主界面布局和主要元素", "提前暴露界面结构不清、操作路径过长的问题", "主界面草图、界面元素清单", "用户能在 1 分钟内理解主流程"],
        ["V4 可交互原型版", "用 HTML 实现可操作界面原型", "验证添加/删除、刷新、图表切换、提醒设置是否顺手", "prototype/index.html、界面原型文档", "核心操作均有反馈，不阻断流程"],
        ["V5 评审修订版", "根据小组评审修正文档和原型", "应对遗漏需求、界面不一致、说明不充分等风险", "迭代评估报告、修订版文档", "checklist 通过，文档与原型一致"],
    ])

    doc.add_heading("4. 本迭代高优先级需求", level=1)
    table(doc, ["优先级", "需求", "本迭代处理方式"], [
        ["Must", "自选股添加/删除", "原型实现输入代码、名称，支持新增和删除。"],
        ["Must", "自选股列表实时刷新", "使用模拟行情刷新现价、涨跌幅、涨跌额，验证界面反馈。"],
        ["Must", "单只股票详情", "主界面展示现价、今开、最高、最低、成交量、成交额。"],
        ["Must", "价格提醒", "原型支持涨跌幅和指定价格提醒规则。"],
        ["Should", "分时图和日 K 趋势", "以折线图方式展示，避免复杂蜡烛图造成进度风险。"],
        ["Could", "持仓备注/标签", "记录为后续扩展，本迭代不实现。"],
    ])

    doc.add_heading("5. 主要任务清单", level=1)
    nums(doc, [
        "制定本迭代计划：明确小组分工、进度安排和滚动式规划。",
        "调研、分析和定义需求：更新 Vision 文档，明确用户、痛点和项目定位。",
        "建立用例模型：用 PlantUML 绘制用例图，将截图放入 SRS，并编写核心用例规约。",
        "设计并实现界面原型：形成主界面和主要元素说明，可用 HTML 实现交互。",
        "小组内部评审：按界面 checklist 检查可用性、美观性、一致性和反馈信息。",
        "编写迭代评估报告：记录评审意见、修订内容、问题和开发总结。"
    ])

    doc.add_heading("6. 风险与应对", level=1)
    table(doc, ["风险", "表现", "应对措施"], [
        ["需求范围膨胀", "想加入真实交易、收益计算、账户登录", "坚持本迭代只验证盯盘和提醒，不做交易。"],
        ["数据源不稳定", "公开行情接口无法保证课堂演示", "原型使用模拟行情；文档中说明真实实现采用 Provider 适配器。"],
        ["界面不清晰", "用户不知道下一步如何操作", "采用三栏布局并在原型文档中说明元素与用例关系。"],
        ["进度不足", "文档和原型无法同时完成", "优先完成 Must 需求，放弃低优先级需求。"],
        ["文档与原型不一致", "SRS 写了但界面没有体现", "评审时逐项对照需求-用例-界面映射表。"],
    ])

    doc.add_heading("7. 完成标准", level=1)
    bullets(doc, [
        "hw1 目录下包含迭代计划、Vision、SRS、界面原型文档、迭代评估报告。",
        "SRS 中包含用例模型说明，PlantUML 源文件在 docs/diagrams 中保存。",
        "界面原型文档说明主要界面、主要元素，以及界面元素与需求/用例的关系。",
        "原型可打开并能演示核心流程。",
        "GitHub 中能追踪本迭代文档、模型和原型代码。"
    ])

    doc.save(OUT / "StockMonitor-hw1迭代计划-v1.1.docx")


def build_ui_prototype_doc():
    doc = Document()
    setup(doc, "界面原型文档")
    history(doc, "新增独立界面原型文档，满足 hw1 成果要求")

    doc.add_heading("1. 原型目标", level=1)
    doc.add_paragraph("本界面原型用于在编码前验证 Stock Monitor 的主要交互方式和信息架构，重点应对需求风险：用户是否能快速管理自选股、查看行情趋势并设置价格提醒。原型可采用 HTML 交互实现，也可作为界面截图提交；本项目采用 HTML/CSS/JavaScript 实现。")

    doc.add_heading("2. 人机交互方式", level=1)
    bullets(doc, [
        "用户通过鼠标点击选择股票、切换图表、删除自选股和保存提醒。",
        "用户通过键盘输入股票代码、股票名称和提醒阈值。",
        "系统通过行情数字变化、图表刷新、提醒列表和文字状态反馈操作结果。",
        "界面语言采用中文，股票代码和技术字段保留通用英文/数字表达。"
    ])

    doc.add_heading("3. 主要界面", level=1)
    table(doc, ["界面", "主要元素", "说明"], [
        ["主界面", "系统名称、刷新状态、自选股列表、添加区、行情详情、图表区、提醒区", "用户进入系统后的第一屏，覆盖核心盯盘流程。"],
        ["自选股区域", "股票代码输入框、名称输入框、添加按钮、股票列表、删除按钮", "支持添加/删除自选股并查看列表行情。"],
        ["行情详情区域", "现价、涨跌幅、今开、最高、最低、成交额", "展示被选中股票的关键指标。"],
        ["图表区域", "分时标签、日 K 标签、折线图画布", "以简化折线图表达短期和日级趋势。"],
        ["提醒区域", "提醒类型、阈值输入、保存按钮、提醒列表", "支持涨超/跌超、涨过/跌破规则。"],
    ])

    doc.add_heading("4. 界面元素与需求/用例映射", level=1)
    table(doc, ["界面元素", "对应需求", "对应用例", "用户反馈"], [
        ["股票代码/名称输入框 + 添加按钮", "FR-01 添加自选股", "UC-01 管理自选股", "新增股票立即显示在列表中。"],
        ["自选股列表", "FR-02 删除自选股；FR-03 实时行情列表", "UC-01、UC-02", "选中项高亮，价格和涨跌幅周期变化。"],
        ["刷新按钮和刷新状态", "FR-03 行情刷新", "UC-02 查看自选股实时行情", "显示最近刷新时间。"],
        ["行情指标卡片", "FR-04 单只股票详情", "UC-03 查看单只股票详情", "指标随选中股票切换。"],
        ["分时/日 K 标签和折线图", "FR-05 分时图；FR-06 日 K 趋势", "UC-04、UC-05", "标签切换后图表颜色和标题变化。"],
        ["提醒表单和提醒列表", "FR-07/FR-08 价格提醒", "UC-06 设置价格提醒", "保存后列表展示规则，触发后高亮提示。"],
    ])

    doc.add_heading("5. 界面 checklist 检查", level=1)
    table(doc, ["检查项", "结论", "说明"], [
        ["人机交互方式是否合理，是否有创新", "通过", "三栏布局贴合盯盘场景，提醒区减少持续盯屏压力。"],
        ["每个界面是否有系统名称或 logo", "通过", "主界面左上角显示 Stock Monitor。"],
        ["界面风格是否一致", "通过", "统一使用蓝绿辅助色、卡片式指标和一致按钮样式。"],
        ["中文还是英文，或两者都支持", "通过", "界面中文为主，股票代码和技术字段保留标准表达。"],
        ["用户是否需要培训，操作是否方便", "通过", "添加、查看、提醒均在一屏完成。"],
        ["是否有在线帮助或提示", "部分通过", "原型内有元素映射说明，正式产品可增加帮助入口。"],
        ["界面大小能否自适应调整", "通过", "CSS 提供桌面和平板/移动宽度响应式布局。"],
        ["布局、色彩、字体大小是否合理", "通过", "指标、列表、图表层次清楚。"],
        ["界面是否简洁明了", "通过", "不提供交易功能，避免干扰核心盯盘任务。"],
        ["用户操作是否有反馈信息", "通过", "刷新时间、选中状态、提醒触发均有反馈。"],
        ["出错信息是否清晰，有无防错措施", "待增强", "原型已有基础默认值，后续实现需补充接口错误提示。"],
        ["是否从用户角度思考，是否越用越好用", "通过", "围绕商科学生和轻量投资者的核心观察任务设计。"],
    ])

    doc.add_heading("6. 原型文件", level=1)
    bullets(doc, [
        "HTML 原型文件：prototype/index.html。",
        "原型使用模拟行情数据，可离线演示添加、删除、刷新、图表切换和提醒设置。",
        "若后续使用 AI 生成界面，应将 prompt 记录在本文件中；当前原型由项目组按需求手工实现。"
    ])

    doc.save(OUT / "StockMonitor-界面原型.docx")


def build_evaluation():
    doc = Document()
    setup(doc, "hw1 迭代评估报告")
    history(doc, "根据界面原型迭代要求补充评审记录、开发总结和修订结果")

    doc.add_heading("1. 迭代总结", level=1)
    doc.add_paragraph("本次 hw1 迭代围绕界面原型展开，重点应对需求风险。项目组先明确轻量股票盯盘的用户和问题，再通过 SRS、用例模型和 HTML 原型验证核心流程，最后按界面 checklist 完成小组评审和修订。")

    doc.add_heading("2. 完成情况", level=1)
    table(doc, ["要求", "完成状态", "说明"], [
        ["本次迭代计划", "完成", "采用滚动式规划，形成 5 版方案。"],
        ["Vision 文档", "完成", "说明项目定位、用户、功能概述和竞争对比。"],
        ["SRS", "完成", "包含功能/非功能需求、用例模型说明和界面映射。"],
        ["用例规约", "完成", "覆盖自选股、行情、详情、图表和提醒核心用例。"],
        ["界面原型", "完成", "HTML 原型可交互，并新增界面原型文档说明。"],
        ["内部评审与改进", "完成", "按 checklist 记录问题并修正文档。"],
    ])

    doc.add_heading("3. 小组评审记录", level=1)
    table(doc, ["评审发现", "影响", "修改措施"], [
        ["原计划没有突出 hw1 是界面原型迭代", "与课程要求对应不够直接", "重写迭代计划，明确目标为应对需求风险。"],
        ["界面原型说明原本分散在 SRS 和 README 中", "不符合“界面原型.docx”提交习惯", "新增独立界面原型文档。"],
        ["部分低优先级功能会扩大范围", "进度风险增加", "将持仓备注列为 Could，真实交易列为 Won't。"],
        ["行情接口存在不可控风险", "演示可能失败", "原型保留模拟行情，正式架构使用 Provider。"],
        ["缺少 checklist 式评审", "无法体现小组评审过程", "在界面原型文档和评估报告中加入检查结果。"],
    ])

    doc.add_heading("4. 需求风险处理结果", level=1)
    table(doc, ["需求风险", "处理结果"], [
        ["用户是否真的需要复杂交易功能", "通过访谈模拟和评审决定不做交易，只做盯盘与提醒。"],
        ["用户是否能理解界面", "通过三栏原型验证，一屏展示主要流程。"],
        ["图表是否过于复杂", "改为分时折线和日 K 收盘价折线。"],
        ["提醒功能是否有价值", "列为 Must，作为降低盯盘压力的核心卖点。"],
        ["公开行情是否可靠", "原型阶段使用模拟数据，设计阶段保留 Provider 替换机制。"],
    ])

    doc.add_heading("5. 开发经验", level=1)
    bullets(doc, [
        "界面原型迭代不应追求功能越多越好，而应优先验证高价值需求。",
        "用例、需求和界面元素之间建立映射，可以减少文档和原型脱节。",
        "公开行情数据适合课堂验证，但必须提前准备模拟数据以保证演示稳定。",
        "小组评审应使用 checklist，避免只凭主观感觉判断界面质量。"
    ])

    doc.add_heading("6. 遗留问题与后续计划", level=1)
    table(doc, ["遗留问题", "后续处理"], [
        ["错误提示仍较简单", "hw2 技术原型中补充接口错误、超时和缓存提示。"],
        ["PlantUML 图目前以源文件形式保存", "提交前可截图后插入 SRS，满足截图粘贴要求。"],
        ["真实行情数据尚未接入", "hw2 设计并验证 MarketDataProvider。"],
        ["尚未实现完整 Windows 客户端", "本迭代只做界面原型，综合实践阶段再实现完整代码。"],
    ])

    doc.save(OUT / "StockMonitor-hw1迭代评估报告-v1.1.docx")


def build_srs_addendum():
    doc = Document()
    setup(doc, "SRS 补充说明：hw1 用例与界面原型映射")
    history(doc, "补充课程截图要求中的用例模型、界面原型放置规则和需求风险说明")

    doc.add_heading("1. 补充目的", level=1)
    doc.add_paragraph("本补充说明用于强化 hw1 的 SRS 交付：SRS 应包含支撑主要需求的用例模型和核心用例说明；界面原型代码、prompt 或截图不放入 SRS，而放入独立的界面原型文档。")

    doc.add_heading("2. 用例模型", level=1)
    doc.add_paragraph("用例图使用 PlantUML 绘制，源文件位于 docs/diagrams/use_case.puml。提交 Word 前可将 PlantUML 渲染结果截图粘贴到 SRS 中。")
    doc.add_paragraph("""@startuml
left to right direction
actor "个人投资者/商科学生" as User
rectangle "Stock Monitor" {
  usecase "管理自选股" as UC1
  usecase "查看实时行情" as UC2
  usecase "查看股票详情" as UC3
  usecase "查看分时图" as UC4
  usecase "查看日K趋势" as UC5
  usecase "设置价格提醒" as UC6
}
User --> UC1
User --> UC2
User --> UC3
User --> UC4
User --> UC5
User --> UC6
@enduml""")

    doc.add_heading("3. 需求风险与原型验证", level=1)
    table(doc, ["需求风险", "原型验证方式", "结论"], [
        ["用户是否需要复杂交易功能", "主界面不提供交易入口", "本阶段聚焦盯盘与提醒。"],
        ["自选股管理是否方便", "左侧固定列表 + 添加输入区", "操作路径短，适合轻量用户。"],
        ["趋势图是否足够表达需求", "分时和日 K 两个标签切换", "无需复杂蜡烛图即可展示趋势。"],
        ["提醒规则是否容易理解", "右侧表单按类型和阈值输入", "符合用户降低盯盘压力的目标。"],
    ])

    doc.add_heading("4. 界面原型引用", level=1)
    bullets(doc, [
        "独立文档：docs/word/StockMonitor-界面原型.docx。",
        "可交互文件：prototype/index.html。",
        "界面原型文档中说明主要界面、主要元素、界面元素与需求/用例关系。"
    ])

    doc.save(OUT / "StockMonitor-SRS补充说明.docx")


if __name__ == "__main__":
    build_iteration_plan()
    build_ui_prototype_doc()
    build_evaluation()
    build_srs_addendum()
    print("updated hw1 documents")
