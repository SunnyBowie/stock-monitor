from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
WORD = ROOT / "docs" / "word"


def paragraph_text(paragraph):
    return "".join(run.text for run in paragraph.runs) or paragraph.text


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def clean(path):
    doc = Document(path)
    markers = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph_text(paragraph).strip() == "附录：PlantUML 用例图截图"
    ]
    if len(markers) <= 1:
        doc.save(path)
        return

    # Keep the last appendix, because it was inserted after the latest diagram refresh.
    for start, end in zip(markers[:-1], markers[1:]):
        for paragraph in doc.paragraphs[start:end]:
            if paragraph._element is not None:
                delete_paragraph(paragraph)
    doc.save(path)


if __name__ == "__main__":
    for name in ["StockMonitor-软件需求规约SRS.docx", "StockMonitor-SRS补充说明.docx"]:
        path = WORD / name
        clean(path)
        print(path)
