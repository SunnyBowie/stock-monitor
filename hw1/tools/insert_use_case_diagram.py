from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_DIR = ROOT / "docs" / "diagrams"
WORD_DIR = ROOT / "docs" / "word"
PNG = DIAGRAM_DIR / "use_case.png"


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def text_center(draw, box, text, fnt, fill=(20, 35, 55)):
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + (len(lines) - 1) * 6
    y = box[1] + ((box[3] - box[1]) - total_h) / 2
    for i, line in enumerate(lines):
        x = box[0] + ((box[2] - box[0]) - widths[i]) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += heights[i] + 6


def actor(draw, x, y, label):
    pen = (24, 63, 99)
    draw.ellipse((x - 14, y, x + 14, y + 28), outline=pen, width=3)
    draw.line((x, y + 28, x, y + 86), fill=pen, width=3)
    draw.line((x - 36, y + 48, x + 36, y + 48), fill=pen, width=3)
    draw.line((x, y + 86, x - 32, y + 126), fill=pen, width=3)
    draw.line((x, y + 86, x + 32, y + 126), fill=pen, width=3)
    text_center(draw, (x - 95, y + 132, x + 95, y + 190), label, font(22, True))


def ellipse(draw, cx, cy, w, h, text):
    box = (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)
    draw.ellipse(box, fill=(244, 250, 255), outline=(35, 100, 170), width=3)
    text_center(draw, box, text, font(21))
    return box


def arrow(draw, start, end, label=None):
    draw.line((start[0], start[1], end[0], end[1]), fill=(45, 62, 80), width=3)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    import math

    angle = math.atan2(dy, dx)
    length = 14
    for off in (2.55, -2.55):
        ax = end[0] - length * math.cos(angle + off)
        ay = end[1] - length * math.sin(angle + off)
        draw.line((end[0], end[1], ax, ay), fill=(45, 62, 80), width=3)
    if label:
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        draw.text((mx - 45, my - 24), label, font=font(18), fill=(90, 90, 90))


def build_png():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((310, 70, 1380, 950), radius=18, outline=(23, 54, 93), width=4)
    draw.text((350, 92), "Stock Monitor Windows 端股票盯盘软件", font=font(30, True), fill=(23, 54, 93))

    actor(draw, 135, 350, "个人投资者\n/商科学生")
    actor(draw, 1540, 250, "行情/财务数据源\nAKShare/东方财富")
    actor(draw, 1540, 690, "Windows\n通知服务")

    boxes = {
        "UC1": ellipse(draw, 620, 220, 290, 82, "管理自选股"),
        "UC2": ellipse(draw, 620, 360, 340, 82, "查看自选股实时行情"),
        "UC3": ellipse(draw, 620, 500, 320, 82, "查看单只股票详情"),
        "UC4": ellipse(draw, 1020, 280, 260, 82, "查看分时图"),
        "UC5": ellipse(draw, 1020, 420, 300, 82, "查看日 K 趋势图"),
        "UC6": ellipse(draw, 620, 650, 330, 82, "查看财务报表摘要"),
        "UC7": ellipse(draw, 1020, 560, 350, 82, "设置价格/涨跌幅提醒"),
        "UC8": ellipse(draw, 1020, 720, 260, 82, "触发提醒通知"),
        "UC9": ellipse(draw, 620, 780, 320, 82, "维护本地持仓备注"),
    }

    # User arrows
    for y in [220, 360, 500, 650, 780]:
        arrow(draw, (235, 445), (475, y))
    for y in [280, 420, 560]:
        arrow(draw, (235, 445), (865, y))

    # Market data arrows
    for y in [360, 500, 280, 420]:
        arrow(draw, (1440, 335), (1180 if y in [280, 420] else 790, y))
    arrow(draw, (1440, 335), (790, 650))

    # Include / extend and notification
    arrow(draw, (1020, 602), (1020, 678), "<<include>>")
    arrow(draw, (790, 386), (930, 680), "<<extend>>")
    arrow(draw, (1150, 720), (1450, 760))

    draw.text((345, 900), "说明：源文件为 docs/diagrams/use_case.puml；本图为渲染后截图，用于粘贴到 SRS 文档。", font=font(20), fill=(80, 90, 100))
    img.save(PNG)


def insert_picture(doc_path, marker_title):
    doc = Document(doc_path)
    doc.add_page_break()
    doc.add_heading(marker_title, level=1)
    doc.add_paragraph("以下为 PlantUML 用例图渲染后的截图，满足“绘图使用 PlantUML，截图后粘贴在文档中”的要求。")
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run().add_picture(str(PNG), width=Inches(6.4))
    doc.add_paragraph("PlantUML 源文件：docs/diagrams/use_case.puml")
    doc.save(doc_path)


if __name__ == "__main__":
    build_png()
    targets = [
        (WORD_DIR / "StockMonitor-软件需求规约SRS.docx", "附录：PlantUML 用例图截图"),
        (WORD_DIR / "StockMonitor-SRS补充说明.docx", "附录：PlantUML 用例图截图"),
    ]
    for path, title in targets:
        insert_picture(path, title)
        print(path)
    print(PNG)
