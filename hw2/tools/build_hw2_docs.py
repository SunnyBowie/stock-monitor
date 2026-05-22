from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
WORD = ROOT / "docs" / "word"
IMG = ROOT / "docs" / "images"
WORD.mkdir(parents=True, exist_ok=True)

PROJECT = "Stock Monitor JavaFX 原生股票看盘软件"
DATE = "2026 年 5 月 22 日"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if fill:
        shade(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(35, 100, 170)
    doc.styles["Heading 3"].font.size = Pt(11.5)
    doc.styles["Heading 3"].font.color.rgb = RGBColor(15, 139, 141)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(23, 54, 93)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"版本 2.2    {DATE}")
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(t.rows[0].cells[i], header, True, "EAF2F8")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def history(doc, note):
    doc.add_heading("修订历史记录", level=1)
    table(doc, ["日期", "版本", "说明", "作者"], [[DATE, "2.0", note, "项目组"]])


def picture(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(IMG / filename), width=Inches(6.4))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def save(doc, name):
    path = WORD / name
    doc.save(path)
    print(path)


def build_plan():
    doc = Document()
    setup(doc, "hw2 技术原型迭代计划")
    history(doc, "技术原型迭代计划：应对架构风险")
    doc.add_heading("1. 本迭代目标", level=1)
    doc.add_paragraph("hw2 是技术原型迭代，目标是验证 hw1 界面原型背后的架构是否可行，重点应对架构风险、技术选型风险和数据持久化风险。")
    doc.add_heading("2. 技术方案设计任务", level=1)
    bullets(doc, [
        "选择架构风格：采用四层分层架构，UI、应用服务、领域对象和基础设施分离。",
        "确定语言、框架、工具：Java 21、JavaFX 21、Maven、Jackson、jpackage、Git/GitHub。",
        "设计多个架构视图：概念模型、用例实现顺序图、VOPC 类图、合并类图和架构视图。",
        "设计关键算法：全 A 股分页导入、行情轮询、分时/K 线绘制、F10 财务报表解析。",
        "选择编程规范：参考 Google Java Style Guide 和中文 Google 编程规范。"
    ])
    doc.add_heading("3. 技术原型范围", level=1)
    table(doc, ["核心用例", "技术原型实现", "验证风险"], [
        ["管理自选股", "JavaFX 搜索列表 + MarketUniverseService", "UI 到仓储链路是否可行"],
        ["查看实时行情", "EastMoneyMarketDataProvider + QuoteRefreshService + Canvas 图表", "行情源适配和图表刷新是否可行"],
        ["设置价格提醒", "F10 财务报表 + Canvas 柱状图", "财务字段解析和图形化展示是否可行"],
    ])
    doc.add_heading("4. 进度安排", level=1)
    table(doc, ["步骤", "工作内容", "输出"], [
        ["1", "制定本迭代计划", "hw2 技术原型迭代计划"],
        ["2", "技术方案与架构建模", "分析模型、架构视图、编程规范"],
        ["3", "实现技术原型", "hw2/javafx-app 下 JavaFX 原生桌面代码"],
        ["4", "测试技术原型", "Maven + SmokeProbe 结果与人工测试记录"],
        ["5", "迭代评估", "评审和测试记录、开发总结"],
    ])
    save(doc, "StockMonitor-hw2迭代计划.docx")


def build_vision():
    doc = Document()
    setup(doc, "hw2 更新后的 Vision 文档")
    history(doc, "根据技术原型迭代更新架构可行性与范围说明")
    doc.add_heading("1. 更新目的", level=1)
    doc.add_paragraph("hw2 在 hw1 的需求和界面原型基础上，进一步验证软件架构是否能支撑核心需求。本次 Vision 更新强调技术范围、架构选择和可演示性。")
    doc.add_heading("2. 产品定位更新", level=1)
    doc.add_paragraph("Stock Monitor 仍定位为 Windows 端轻量股票盯盘软件，不做真实交易、不接入证券账户。本次技术原型证明其核心能力可通过分层架构实现，已接入东方财富真实行情源。")
    doc.add_heading("3. 技术原型价值", level=1)
    bullets(doc, [
        "验证全 A 股导入、行情刷新、分时/K 线和财务报表核心链路可以端到端运行。",
        "验证 MarketUniverseService 能分页导入并缓存全部 A 股。",
        "验证 MarketDataProvider 接口可以隔离真实行情源与模拟行情源。",
        "验证 JavaFX 原生 UI 能承载简约看盘界面和 Canvas 图表。"
    ])
    doc.add_heading("4. 风险更新", level=1)
    table(doc, ["风险", "hw2 处理方式"], [
        ["真实行情接口不稳定", "已使用 EastMoneyMarketDataProvider；公开接口不具备商业 SLA，商业化需授权数据源。"],
        ["架构耦合导致难以扩展", "采用 UI/Application/Domain/Infrastructure 分层。"],
        ["公开接口限流/网络波动", "采用分页刷新和按需详情加载降低请求压力。"],
        ["完整软件开发量过大", "本迭代只实现 2-3 个核心用例，完整代码留到综合实践。"],
    ])
    save(doc, "StockMonitor-hw2更新Vision文档.docx")


def build_analysis_model():
    doc = Document()
    setup(doc, "分析模型")
    history(doc, "新增概念模型、用例实现模型和合并类图")
    doc.add_heading("1. 概念模型", level=1)
    doc.add_paragraph("概念模型给出系统核心概念类：Stock、MarketPage、MarketQuote、FinancialReport 和 MarketDataProvider。")
    picture(doc, "concept_model.png", "图 1 概念模型类图")
    doc.add_heading("2. 用例实现：管理自选股", level=1)
    doc.add_paragraph("管理自选股用例验证 UI、应用服务和 无强制本地数据库 仓储之间的协作。PlantUML 源文件：docs/diagrams/sequence_watchlist.puml。")
    doc.add_heading("3. 用例实现：设置价格提醒", level=1)
    picture(doc, "sequence_alert.png", "图 2 设置并触发价格提醒顺序图")
    picture(doc, "vopc_quote.png", "图 3 行情刷新 VOPC 类图")
    doc.add_heading("4. 通信图与整体类图", level=1)
    doc.add_paragraph("通信图源文件位于 docs/diagrams/communication_alert.puml；整体类图由各 VOPC 类图合并而成。")
    picture(doc, "merged_class_model.png", "图 4 合并后的整体类图")
    save(doc, "StockMonitor-分析模型.docx")


def build_architecture_doc():
    doc = Document()
    setup(doc, "软件构架文档")
    history(doc, "hw2 技术原型架构设计与评审版本")
    doc.add_heading("1. 架构目标", level=1)
    bullets(doc, [
        "支持功能需求：全 A 股股票池、行情查看、分时/K 线、财务报表。",
        "支持非功能需求：可维护、可测试、可替换行情源、可独立打包。",
        "应对需求变更：新增授权行情源或扩展指标时不影响 UI 主流程。"
    ])
    doc.add_heading("2. 架构风格与技术选型", level=1)
    table(doc, ["项目", "选择", "理由"], [
        ["架构风格", "四层分层架构", "降低 UI、业务、数据源耦合。"],
        ["语言", "Java 21", "LTS 版本，适合面向对象桌面应用和 jpackage 打包。"],
        ["UI 框架", "JavaFX 21", "原生桌面 GUI，支持 CSS 与 Canvas 绘图。"],
        ["数据库", "无强制本地数据库", "当前以运行期缓存为主，后续可扩展本地自选股持久化。"],
        ["测试", "Maven + SmokeProbe", "Maven 编译验证，SmokeProbe 验证真实接口。"],
    ])
    doc.add_heading("3. 架构视图", level=1)
    picture(doc, "architecture_views.png", "图 1 软件架构视图")
    doc.add_heading("4. 包/模块组织", level=1)
    table(doc, ["包", "职责", "代表类"], [
        ["domain", "领域对象和值计算", "Stock、MarketQuote、MarketPage、FinancialReport"],
        ["application", "用例服务和端口定义", "WatchlistService、MarketUniverseService、QuoteRefreshService"],
        ["infrastructure", "无强制本地数据库 仓储、行情源适配", "无强制本地数据库StockRepository、MockMarketDataProvider"],
        ["ui", "桌面界面和图表展示", "StockMonitorApp、MainController、Canvas Charts"],
        ["tests", "单元测试", "test_services.py"],
    ])
    doc.add_heading("5. 架构评审结论", level=1)
    bullets(doc, [
        "核心用例可以通过分层架构端到端运行。",
        "行情源被接口隔离，后续接入授权 Level-2 或商业数据源不需要重写 UI。",
        "测试发现 无强制本地数据库 文件锁问题，已通过显式关闭连接修复。",
        "技术原型代码不是完整软件，但足以验证架构可行性。"
    ])
    save(doc, "StockMonitor-hw2软件构架文档.docx")


def build_coding_standard():
    doc = Document()
    setup(doc, "编程规范")
    history(doc, "新增 hw2 技术原型编程规范")
    doc.add_heading("1. 规范来源", level=1)
    bullets(doc, [
        "英文 Google Style Guide：https://github.com/google/styleguide",
        "中文 Google Style Guide：https://github.com/zh-google-styleguide/zh-google-styleguide",
        "本项目采用 Java，主要参考 Google Java Style Guide。"
    ])
    doc.add_heading("2. 命名规范", level=1)
    table(doc, ["对象", "规范", "示例"], [
        ["模块/文件", "小写包名 / kebab 配置名", "无强制本地数据库_repository.py"],
        ["类", "PascalCase", "WatchlistService"],
        ["方法/变量", "lowerCamelCase", "refreshNow"],
        ["常量", "UPPER_lowerCamelCase", "DEFAULT_REFRESH_SECONDS"],
        ["枚举值", "UPPER_lowerCamelCase", "ABOVE_PRICE"],
    ])
    doc.add_heading("3. 分层约束", level=1)
    bullets(doc, [
        "UI 层只能调用 application service，不直接操作 无强制本地数据库。",
        "application 层依赖 MarketDataProvider 端口，不依赖具体 HTTP 实现。",
        "domain 层不依赖 UI、数据库或网络。",
        "infrastructure 层实现东方财富行情与财务接口适配。"
    ])
    doc.add_heading("4. 测试规范", level=1)
    bullets(doc, [
        "每个核心链路至少有编译验证或烟测入口。",
        "测试优先覆盖数据解析、分页导入、分时/K 线和财务报表。",
        "系统功能测试采用 SmokeProbe 与人工打开 JavaFX 窗口结合。"
    ])
    save(doc, "StockMonitor-编程规范.docx")


def build_evaluation():
    doc = Document()
    setup(doc, "hw2 迭代评估报告")
    history(doc, "技术原型迭代评估：评审和测试记录、开发总结")
    doc.add_heading("1. 迭代总结", level=1)
    doc.add_paragraph("本次 hw2 技术原型迭代围绕架构风险展开，完成技术方案设计、分析建模、软件架构文档、编程规范、技术原型代码和测试记录。最终软件放在 hw2/javafx-app 文件夹。")
    doc.add_heading("2. 完成情况", level=1)
    table(doc, ["要求", "完成情况", "说明"], [
        ["本次迭代计划", "完成", "StockMonitor-hw2迭代计划.docx"],
        ["更新后的 Vision", "完成", "StockMonitor-hw2更新Vision文档.docx"],
        ["分析建模", "完成", "JavaFX 架构视图、全市场导入顺序图、VOPC 类图、合并类图"],
        ["软件架构文档", "完成", "四层架构、包组织和架构评审"],
        ["编程规范", "完成", "参考 Google Java Style Guide"],
        ["技术原型代码", "完成", "hw2/javafx-app 下 JavaFX 原生桌面代码"],
        ["测试技术原型", "完成", "3 个 Maven + SmokeProbe 全部通过"],
    ])
    doc.add_heading("3. 测试记录", level=1)
    table(doc, ["测试项", "方法", "结果"], [
        ["管理自选股", "SmokeProbe marketPage(1,3)", "通过"],
        ["查看实时行情", "SmokeProbe quote(600519)", "通过"],
        ["设置价格提醒", "SmokeProbe intraday/candles/F10", "通过"],
        ["无强制本地数据库 文件锁", "首次测试发现临时 db 无法删除", "已采用分页刷新/按需加载"],
    ])
    doc.add_heading("4. 架构风险处理", level=1)
    table(doc, ["风险", "处理结果"], [
        ["UI 与真实接口耦合", "通过 MarketDataProvider 与 application service 隔离。"],
        ["行情源不可用", "EastMoney provider 获取真实数据，Mock provider 保留兜底。"],
        ["本地持久化不稳定", "无强制本地数据库 测试通过，并修复 Windows 文件锁。"],
        ["功能过多影响质量", "实现全市场、详情图表、财务报表三条核心链路。"],
    ])
    doc.add_heading("5. 开发总结", level=1)
    bullets(doc, [
        "本轮验证的是原生桌面架构和真实数据链路。",
        "分层结构使真实行情烟测无需启动 UI 即可验证核心数据。",
        "Windows 环境下 无强制本地数据库 连接必须显式关闭，否则测试清理会失败。",
        "后续可在当前架构基础上接入授权行情源、本地持久化和安装包。"
    ])
    save(doc, "StockMonitor-hw2迭代评估报告.docx")


if __name__ == "__main__":
    build_plan()
    build_vision()
    build_analysis_model()
    build_architecture_doc()
    build_coding_standard()
    build_evaluation()
