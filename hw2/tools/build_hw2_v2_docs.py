from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
WORD = ROOT / "docs" / "word"
DATE = "2026 年 5 月 22 日"
PROJECT = "Stock Monitor JavaFX 原生股票看盘软件"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(35, 100, 170)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    r.bold = True
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"版本 2.1    {DATE}")
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True, "EAF2F8")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def save(doc, name):
    path = WORD / name
    doc.save(path)
    print(path)


def build():
    doc = Document()
    setup(doc, "hw2 第二版技术原型迭代报告")
    doc.add_heading("修订历史记录", level=1)
    table(doc, ["日期", "版本", "说明", "作者"], [[DATE, "2.2", "升级为 JavaFX 原生桌面端、全 A 股导入、真实分时/K线/F10财务数据、jpackage 打包", "项目组"]])

    doc.add_heading("1. 第二版迭代目标", level=1)
    bullets(doc, [
        "将软件完全整合为 JavaFX 原生桌面程序，不依赖外部浏览器打开。",
        "启动后导入全部 A 股股票池，当前东方财富批量行情接口返回 total=5530。",
        "后台采用分页批量刷新机制持续更新全市场行情，详情页按需刷新单股快照、分时、K 线和财务报表。",
        "保留简约界面原则：左侧全市场搜索/自选，右侧单股详情、图表和财务摘要。"
    ])

    doc.add_heading("2. 技术实现", level=1)
    table(doc, ["模块", "实现", "说明"], [
        ["桌面界面", "JavaFX 21 + CSS", "纯原生窗口、暗色简约视觉、ListView、ScrollPane、Canvas 图表。"],
        ["全 A 股导入", "MarketUniverseService", "调用 clist 批量行情接口，200 支/页导入并缓存全市场股票池。"],
        ["持续刷新", "ScheduledExecutorService", "后台轮询分页刷新，避免对公开接口做 5530 支逐只高频请求。"],
        ["真实行情适配", "EastMoneyMarketDataProvider", "接入 stock/get、clist/get、trends2/get、kline/get。"],
        ["真实财务", "EastMoney Datacenter F10", "接入利润表、资产负债表、现金流量表，绘制近四期营收和归母净利润。"],
        ["独立打包", "Maven + jpackage", "将 JavaFX 依赖和运行时打包为 Windows app-image。"],
    ])

    doc.add_heading("3. 验证记录", level=1)
    table(doc, ["验证项", "结果"], [
        ["Maven 构建", "mvn -DskipTests package 成功，编译 21 个 Java 主源码文件。"],
        ["全市场接口", "marketPage(1,3) 返回 total=5530，证明可导入全部 A 股。"],
        ["单股行情", "600519 贵州茅台返回 price=1290.20、change=-1.59%、quoteTime=2026-05-22T16:11:47。"],
        ["分时/K 线", "600519 分时点数 241，K 线数据 575 条。"],
        ["财务报表", "600519 F10 利润表返回营收 4 期、归母净利润 4 期。"],
    ])

    doc.add_heading("4. 当前限制与后续改进", level=1)
    bullets(doc, [
        "东方财富公开接口不是商业 SLA，商业化版本仍需接入授权行情源和 Level-2 数据服务。",
        "全市场行情采用批量分页刷新；详情页分时、K 线和财务报表按需加载，避免 UI 卡顿与接口过载。",
        "A 股休市时仍展示最近一次交易快照，后台刷新频率可按交易时间进一步动态降低。",
        "当前打包目标为 Windows app-image，后续可补充 MSI 安装包、自动升级和本地持久化自选股。"
    ])

    doc.add_heading("5. 运行方式", level=1)
    doc.add_paragraph("开发运行：进入 hw2/javafx-app，执行 ..\\tools\\apache-maven-3.9.11\\bin\\mvn.cmd javafx:run。")
    doc.add_paragraph("构建验证：执行 ..\\tools\\apache-maven-3.9.11\\bin\\mvn.cmd -DskipTests package。")
    doc.add_paragraph("真实数据烟测：执行 D:\\java21\\bin\\java.exe --module-path \"target\\stock-monitor-javafx-0.4.0.jar;target\\dependency\" --module com.stockmonitor/com.stockmonitor.tool.EastMoneySmokeProbe。")
    doc.add_paragraph("独立打包：执行 .\\packaging\\package-windows.ps1，产物输出到 hw2/javafx-app/target/installer。")
    save(doc, "StockMonitor-hw2第二版迭代报告.docx")


if __name__ == "__main__":
    build()
