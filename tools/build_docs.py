from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "word"
OUT.mkdir(parents=True, exist_ok=True)

PROJECT = "Stock Monitor Windows 端股票盯盘软件"
DATE = "2026 年 5 月 21 日"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 20, "17365D"),
        ("Heading 1", 15, "17365D"),
        ("Heading 2", 12.5, "2364AA"),
        ("Heading 3", 11.5, "0F8B8D"),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(23, 54, 93)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"版本 1.0    {DATE}")
    doc.add_paragraph()


def add_history(doc):
    doc.add_heading("修订历史记录", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["日期", "版本", "说明", "作者"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for i, v in enumerate([DATE, "1.0", "课程项目初稿：需求、架构、原型与迭代交付", "项目组"]):
        set_cell_text(table.rows[1].cells[i], v)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "EAF2F8")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def save(doc, name):
    path = OUT / name
    doc.save(path)
    return path


def build_vision():
    doc = Document()
    style_doc(doc, "前景文档")
    add_history(doc)
    doc.add_heading("1. 简介", level=1)
    doc.add_heading("1.1 目的", level=2)
    doc.add_paragraph("本文档用于收集、分析并定义 Stock Monitor 的高层次需求和产品愿景，说明项目为什么值得做、为谁做、解决什么问题，以及第一阶段应交付哪些核心能力。")
    doc.add_heading("1.2 范围", level=2)
    doc.add_paragraph("系统范围限定为 Windows 端轻量股票盯盘工具，不包含真实交易下单、账户登录、证券账户资金划转、投资建议生成和合规投顾服务。")
    doc.add_heading("1.3 定义、缩略语", level=2)
    add_table(doc, ["术语", "说明"], [
        ["自选股", "用户主动关注的股票列表，类似商科资产/持仓观察清单。"],
        ["分时图", "交易日内按照时间采样的价格折线。"],
        ["日 K", "以交易日为单位的历史行情；本项目只展示收盘价趋势线。"],
        ["行情源", "提供股票报价、成交量、成交额和历史价格的数据接口。"],
    ])
    doc.add_heading("1.4 参考资料", level=2)
    bullets(doc, [
        "《se00.平时作业.pdf》课程项目要求。",
        "AKShare GitHub 与文档：开源财经数据接口库，支持 A 股实时行情与历史数据。",
        "BaoStock PyPI/官方资料：用于中国股票市场历史数据获取。",
        "项目 Word 模板：前景文档、用例规约、软件构架文档、SRS、迭代计划、迭代评估报告。"
    ])

    doc.add_heading("2. 定位", level=1)
    doc.add_heading("2.1 商机", level=2)
    doc.add_paragraph("个人投资者、金融/商科学生和课程项目团队常常需要同时观察多只股票的价格变化，但专业交易软件功能复杂、账户绑定成本高，网页行情又不适合长期盯盘。Stock Monitor 以低门槛、可解释、可演示为目标，提供一套小而完整的 Windows 端盯盘体验。")
    doc.add_heading("2.2 问题说明", level=2)
    add_table(doc, ["问题", "影响", "成功标准"], [
        ["多只股票价格变化分散在不同页面", "用户需要频繁切换网页，遗漏关键波动", "自选股列表集中显示现价、涨跌幅、涨跌额"],
        ["商科用户只需要趋势判断，不需要复杂交易终端", "学习成本过高，课堂展示难聚焦", "用简化分时图和收盘价折线表达趋势"],
        ["用户不能长期盯屏", "容易错过涨跌触发点", "本地提醒规则能在刷新时触发文字/弹窗提示"],
    ])
    doc.add_heading("2.3 产品定位说明", level=2)
    doc.add_paragraph("对于需要低门槛观察 A 股行情的个人用户和课程项目团队，Stock Monitor 是一款 Windows 端轻量股票盯盘软件。它区别于东方财富、同花顺等全功能证券终端，聚焦自选股管理、实时行情摘要、趋势图和本地价格提醒，便于课程实现、演示和后续扩展。")

    doc.add_heading("3. 涉众和用户说明", level=1)
    add_table(doc, ["名称", "角色", "关注点"], [
        ["个人投资者/商科学生", "主要用户", "快速查看关注股票走势，获得价格提醒。"],
        ["课程教师/助教", "验收方", "项目是否满足 OOAD 文档、原型和迭代要求。"],
        ["开发小组", "实施方", "数据源可用性、Windows 技术栈、进度和质量。"],
        ["行情数据源提供方", "外部系统", "接口稳定性、调用频率、授权边界。"],
    ])
    doc.add_heading("3.1 用户环境", level=2)
    doc.add_paragraph("用户使用 Windows 10/11 电脑，网络可访问公开行情源。典型使用场景为课堂演示、个人学习或轻量观察，不要求证券账户登录。")
    doc.add_heading("3.2 关键用户需要", level=2)
    add_table(doc, ["需要", "优先级", "当前方案", "提议方案"], [
        ["集中管理自选股", "高", "浏览器收藏或证券软件自选", "本地自选股列表，支持添加/删除。"],
        ["快速看实时价格", "高", "多个网页或手机 App", "列表自动刷新现价、涨跌幅、涨跌额。"],
        ["理解趋势", "中", "专业 K 线图复杂", "简化分时图和日收盘价折线。"],
        ["减少盯盘压力", "中", "人工盯屏", "涨跌幅和价格阈值提醒。"],
    ])

    doc.add_heading("4. 产品概述", level=1)
    doc.add_heading("4.1 产品总体效果", level=2)
    doc.add_paragraph("系统启动后显示三栏布局：左侧自选股列表，中间为单只股票行情详情和图表，右侧为提醒规则与触发提示。用户无需登录证券账户即可完成观察、分析和提醒配置。")
    doc.add_heading("4.2 功能摘要", level=2)
    bullets(doc, [
        "自选股管理：添加、删除股票代码和名称，保存到本地。",
        "实时行情查看：自选股列表周期刷新现价、涨跌幅、涨跌额。",
        "股票详情：展示现价、今开、最高、最低、成交量、成交额。",
        "图表展示：分时折线图和日 K 收盘价趋势线。",
        "价格提醒：支持涨超/跌超百分比、涨过/跌破指定价格。"
    ])
    doc.add_heading("4.3 竞争对比", level=2)
    add_table(doc, ["产品", "优势", "不足", "本项目差异"], [
        ["东方财富/同花顺", "行情全面，功能成熟", "功能重、学习成本高", "只做课程所需的轻量盯盘闭环。"],
        ["网页财经行情", "免安装、数据丰富", "多股集中盯盘和提醒弱", "Windows 端集中列表与本地提醒。"],
        ["Excel/手工记录", "熟悉、可自定义", "实时性差", "自动刷新和图形化展示。"],
    ])
    doc.add_heading("5. 约束与假设", level=1)
    bullets(doc, [
        "课程版本不提供真实交易功能，避免金融合规风险。",
        "公开行情接口可能不稳定，必须设计 mock 数据和适配器切换机制。",
        "刷新频率控制在 5 秒以上，避免对公开接口造成高频压力。",
        "商业化发布需要采购授权行情数据。"
    ])
    return save(doc, "StockMonitor-前景文档.docx")


def build_srs():
    doc = Document()
    style_doc(doc, "软件需求规约 SRS")
    add_history(doc)
    doc.add_heading("1. 简介", level=1)
    doc.add_paragraph("本文档定义 Stock Monitor 的外部行为、功能性需求、非功能性需求、接口与设计约束，为设计、实现和测试提供依据。")
    doc.add_heading("2. 整体说明", level=1)
    doc.add_paragraph("Stock Monitor 是 Windows 桌面端股票盯盘软件，目标用户为个人投资者、商科学生和课程项目评审人员。系统以本地自选股和提醒规则为核心，通过行情数据源适配器获取实时与历史数据。")
    doc.add_heading("2.1 产品功能", level=2)
    bullets(doc, [
        "管理自选股：股票代码、名称、删除和持久化。",
        "实时刷新行情：现价、涨跌幅、涨跌额。",
        "查看详情：今开、最高、最低、成交量、成交额。",
        "查看图表：简化分时折线和日 K 收盘价折线。",
        "价格提醒：按百分比或指定价格触发本地提示。"
    ])
    doc.add_heading("2.2 数据源可获得性", level=2)
    doc.add_paragraph("调研结果表明，课程项目所需字段可通过 AKShare、BaoStock、东方财富/新浪/腾讯公开接口组合获得。实时盯盘优先 AKShare 东方财富实时行情或新浪/腾讯接口，历史日线优先 AKShare/BaoStock。公开接口没有商业 SLA，系统必须支持模拟数据和接口替换。")
    doc.add_heading("3. 具体需求", level=1)
    doc.add_heading("3.1 用例图", level=2)
    doc.add_paragraph("PlantUML 源文件位于 docs/diagrams/use_case.puml。核心参与者为个人投资者/商科学生、行情数据源和 Windows 通知服务。")
    doc.add_paragraph("""@startuml
个人投资者 --> (管理自选股)
个人投资者 --> (查看实时行情)
个人投资者 --> (查看股票详情)
个人投资者 --> (设置价格提醒)
行情数据源 --> (查看实时行情)
(查看实时行情) ..> (触发提醒通知) : <<extend>>
@enduml""")
    doc.add_heading("3.2 功能性需求", level=2)
    add_table(doc, ["编号", "需求", "优先级", "验收标准"], [
        ["FR-01", "用户可以添加股票代码和名称到自选股列表。", "高", "输入合法代码和名称后列表出现新股票并持久保存。"],
        ["FR-02", "用户可以删除自选股。", "高", "删除后该股票不再显示，重启后仍被删除。"],
        ["FR-03", "系统周期刷新自选股现价、涨跌幅、涨跌额。", "高", "刷新后列表字段更新，并显示刷新时间。"],
        ["FR-04", "用户选择单只股票后可查看现价、今开、最高、最低、成交量、成交额。", "高", "详情区域显示 6 个指标且与选中股票一致。"],
        ["FR-05", "系统显示简化分时图。", "中", "以折线方式展示交易日内价格采样。"],
        ["FR-06", "系统显示日 K 收盘价趋势线。", "中", "以折线方式展示最近若干交易日收盘价。"],
        ["FR-07", "用户可以设置涨超/跌超百分比提醒。", "高", "行情刷新满足阈值时出现提醒。"],
        ["FR-08", "用户可以设置涨过/跌破指定价格提醒。", "高", "行情刷新满足价格条件时出现提醒。"],
    ])
    doc.add_heading("3.3 非功能性需求", level=2)
    add_table(doc, ["类别", "需求"], [
        ["易用性", "主要任务不超过 3 次点击；界面字段使用商科用户熟悉的行情术语。"],
        ["可靠性", "行情源失败时显示最后一次缓存和错误提示，不导致程序崩溃。"],
        ["性能", "自选股 50 只以内时，单次刷新响应时间不超过 2 秒；图表切换不超过 500ms。"],
        ["可支持性", "行情源通过 Provider 接口扩展；日志记录接口错误和提醒触发事件。"],
        ["可移植性", "目标运行环境为 Windows 10/11。"],
        ["安全性", "不采集证券账户、交易密码和身份证等敏感信息。"],
    ])
    doc.add_heading("3.4 接口需求", level=2)
    add_table(doc, ["接口", "说明"], [
        ["用户界面", "三栏桌面布局：自选股、详情图表、提醒面板。"],
        ["软件接口", "MarketDataProvider 提供 getQuote、getIntradaySeries、getDailyCloseSeries。"],
        ["数据接口", "SQLite 保存自选股、提醒规则和缓存行情。"],
        ["通信接口", "HTTP/HTTPS 访问行情源，需设置超时、重试和降级策略。"],
    ])
    doc.add_heading("3.5 界面原型与需求映射", level=2)
    add_table(doc, ["界面元素", "对应需求/用例"], [
        ["左侧自选股列表和添加区", "FR-01、FR-02、UC-01 管理自选股"],
        ["列表价格、涨跌幅字段", "FR-03、UC-02 查看实时行情"],
        ["中部行情指标卡片", "FR-04、UC-03 查看股票详情"],
        ["分时/日 K 标签和折线图", "FR-05、FR-06、UC-04/UC-05 查看趋势"],
        ["右侧提醒表单与提醒列表", "FR-07、FR-08、UC-06 设置价格提醒"],
    ])
    return save(doc, "StockMonitor-软件需求规约SRS.docx")


def build_use_cases():
    doc = Document()
    style_doc(doc, "用例规约")
    add_history(doc)
    use_cases = [
        ("UC-01 管理自选股", "用户维护关注股票清单，为后续行情刷新和提醒提供对象。",
         ["用户在添加区输入股票代码和名称。", "系统校验代码和名称非空且代码未重复。", "系统保存股票到本地自选股列表。", "系统刷新列表并默认选中新股票。"],
         ["代码重复：系统提示该股票已在自选股中。", "用户删除股票：系统确认后删除并清理相关提醒。"],
         "应用已启动。", "自选股列表与本地存储保持一致。"),
        ("UC-02 查看自选股实时行情", "用户在自选股列表中观察多只股票的现价、涨跌额和涨跌幅。",
         ["系统按设定周期请求行情源。", "行情源返回自选股报价。", "系统计算涨跌额、涨跌幅并更新列表。", "系统记录刷新时间。"],
         ["行情源失败：系统显示上次缓存数据和错误提示。", "网络不可用：系统切换模拟数据或离线状态。"],
         "至少存在一只自选股。", "列表展示最新可用行情或明确的离线状态。"),
        ("UC-03 查看单只股票详情", "用户选择某只股票并查看关键行情指标。",
         ["用户点击自选股列表中的股票。", "系统加载该股票详情。", "系统显示现价、今开、最高、最低、成交量、成交额。"],
         ["选中股票数据缺失：系统显示占位并提示稍后刷新。"],
         "自选股列表已加载。", "详情区展示选中股票的行情指标。"),
        ("UC-04 查看分时图", "用户通过简化折线观察交易日内价格趋势。",
         ["用户选择分时标签。", "系统获取或追加当日采样点。", "系统绘制分时价格折线。"],
         ["实时分钟数据不可用：系统用轮询采样形成简化分时。"],
         "已选中股票。", "图表区显示分时折线。"),
        ("UC-05 查看日 K 趋势", "用户通过收盘价折线观察中短期趋势。",
         ["用户选择日 K 标签。", "系统获取最近交易日收盘价。", "系统绘制收盘价趋势线。"],
         ["历史数据源失败：系统提示无法加载历史行情。"],
         "已选中股票。", "图表区显示日收盘价折线。"),
        ("UC-06 设置价格提醒", "用户设置涨跌幅或价格阈值，降低人工盯盘压力。",
         ["用户选择提醒类型。", "用户输入阈值。", "系统校验阈值合法。", "系统保存提醒规则。", "行情刷新时系统评估规则，满足条件则提示。"],
         ["阈值非法：系统提示重新输入。", "提醒已触发：系统标记触发状态，避免连续重复弹窗。"],
         "已选中股票。", "提醒规则保存，并在满足条件时出现文字或弹窗提示。"),
    ]
    for name, desc, basic, alt, pre, post in use_cases:
        doc.add_heading(name, level=1)
        doc.add_heading("简要说明", level=2)
        doc.add_paragraph(desc)
        doc.add_heading("基本流", level=2)
        numbered(doc, basic)
        doc.add_heading("备选流", level=2)
        bullets(doc, alt)
        doc.add_heading("特殊需求", level=2)
        doc.add_paragraph("刷新、图表绘制和提醒判断不得阻塞主界面；所有行情接口异常均需可见提示。")
        doc.add_heading("前置条件", level=2)
        doc.add_paragraph(pre)
        doc.add_heading("后置条件", level=2)
        doc.add_paragraph(post)
    return save(doc, "StockMonitor-用例规约.docx")


def build_architecture():
    doc = Document()
    style_doc(doc, "软件构架文档")
    add_history(doc)
    doc.add_heading("1. 简介", level=1)
    doc.add_paragraph("本文档从构架视角说明 Stock Monitor 的关键设计决策、逻辑分层、进程视图、部署视图和质量属性。")
    doc.add_heading("2. 构架表示方式", level=1)
    doc.add_paragraph("采用 4+1 视图描述：用例视图、逻辑视图、进程视图、部署视图、实施视图，以及数据视图。PlantUML 架构图位于 docs/diagrams/architecture.puml。")
    doc.add_heading("3. 构架目标和约束", level=1)
    bullets(doc, [
        "面向对象设计，核心业务对象与 UI 解耦。",
        "行情源不稳定，必须使用适配器和降级策略。",
        "Windows 桌面端优先，技术栈可选 WPF/WinUI；课程原型用 HTML 表达界面。",
        "不实现交易下单，降低安全和合规复杂度。"
    ])
    doc.add_heading("4. 用例视图", level=1)
    doc.add_paragraph("构架重点支撑 UC-01 管理自选股、UC-02 查看实时行情、UC-03 查看详情、UC-04/05 图表展示和 UC-06 价格提醒。")
    doc.add_heading("5. 逻辑视图", level=1)
    add_table(doc, ["层", "职责", "主要类/组件"], [
        ["表示层", "页面、控件、用户交互、图表显示", "MainWindow、WatchlistView、QuoteDetailView、AlertPanel"],
        ["应用层", "编排用例流程、刷新调度、提醒判断", "WatchlistService、QuoteRefreshService、AlertService"],
        ["领域层", "表达股票、行情、提醒规则等核心对象", "Stock、MarketQuote、IntradayPoint、DailyClosePoint、AlertRule"],
        ["基础设施层", "外部接口、本地存储、通知实现", "AkShareProvider、SinaProvider、SQLiteRepository、WindowsNotificationAdapter"],
    ])
    doc.add_heading("6. 进程视图", level=1)
    doc.add_paragraph("主 UI 线程负责界面响应；后台定时任务负责行情轮询；提醒服务在每次刷新后同步评估规则。网络请求需要超时控制，结果通过事件或观察者模式通知 UI。")
    doc.add_heading("7. 部署视图", level=1)
    doc.add_paragraph("系统部署在用户 Windows 10/11 电脑。应用本体、本地 SQLite 数据库和配置文件位于用户目录；行情请求通过 HTTPS 访问外部数据源；提醒通过 Windows Toast 或应用内弹窗显示。")
    doc.add_heading("8. 实施视图", level=1)
    doc.add_paragraph("建议 GitHub 仓库按 `src/`、`tests/`、`docs/`、`prototype/` 组织。实现阶段采用分支开发和 Pull Request 评审，文档与原型同步提交。")
    doc.add_heading("9. 数据视图", level=1)
    add_table(doc, ["实体", "关键字段", "说明"], [
        ["Stock", "code, name, market", "自选股基础信息。"],
        ["MarketQuote", "price, open, high, low, volume, amount, change, pct", "实时行情快照。"],
        ["AlertRule", "stockCode, ruleType, threshold, enabled, triggeredAt", "提醒规则。"],
        ["QuoteCache", "stockCode, sampledAt, payload", "最近行情缓存与分时采样。"],
    ])
    doc.add_heading("10. 大小和性能", level=1)
    bullets(doc, [
        "课程版本支持 50 只以内自选股。",
        "默认刷新周期 8 到 15 秒，可配置但不低于 5 秒。",
        "图表点数控制在 240 个以内，避免 UI 卡顿。"
    ])
    doc.add_heading("11. 质量", level=1)
    bullets(doc, [
        "可扩展性：新增行情源只需实现 Provider 接口。",
        "可靠性：接口失败不影响本地自选股和提醒规则。",
        "可测试性：领域层和应用层可用模拟 Provider 单元测试。",
        "可维护性：UI、业务、数据访问分层明确。"
    ])
    return save(doc, "StockMonitor-软件构架文档.docx")


def build_iteration_plan():
    doc = Document()
    style_doc(doc, "迭代计划")
    add_history(doc)
    doc.add_heading("1. 迭代目标", level=1)
    doc.add_paragraph("采用 3 次迭代完成课程项目：先完成需求和原型，再实现核心盯盘闭环，最后补齐提醒、文档评审和演示材料。")
    add_table(doc, ["迭代", "时间", "目标", "主要产出"], [
        ["迭代 1：调研与需求", "第 1 周", "确认数据源、用户和核心需求", "前景文档、SRS 初稿、用例图、界面草图"],
        ["迭代 2：核心原型", "第 2 周", "完成自选股、行情详情、图表原型", "HTML 原型、架构文档、用例规约"],
        ["迭代 3：提醒与评审", "第 3 周", "完成提醒逻辑、文档修订、演示准备", "迭代评估报告、最终 Word 文档、GitHub 仓库整理"],
    ])
    doc.add_heading("2. 任务分解", level=1)
    add_table(doc, ["任务", "负责人", "优先级", "验收标准"], [
        ["数据源可获得性分析", "需求/后端", "高", "明确实时、历史、提醒所需字段来源和风险。"],
        ["Vision 与 SRS", "需求", "高", "覆盖用户、问题、功能和非功能需求。"],
        ["用例图与用例规约", "需求/设计", "高", "至少覆盖自选股、行情、图表、提醒。"],
        ["HTML 界面原型", "前端", "高", "可打开、可交互、能展示主界面元素。"],
        ["架构设计", "后端/架构", "中", "分层、核心类、数据视图和部署视图清晰。"],
        ["小组评审与修改", "全组", "高", "记录问题、改文档、改原型。"],
    ])
    doc.add_heading("3. 风险与应对", level=1)
    add_table(doc, ["风险", "概率", "影响", "应对"], [
        ["公开行情接口变更或不可用", "中", "高", "使用 Provider 适配器和 mock 数据。"],
        ["图表实现超出课程时间", "中", "中", "只做折线图，不做复杂蜡烛图。"],
        ["文档模板内容多导致进度拖延", "中", "中", "聚焦课程要求章节，先完成可评审版本。"],
        ["组员 Git 协作冲突", "中", "中", "按目录分工，PR 评审后合并。"],
    ])
    doc.add_heading("4. GitHub 管理方案", level=1)
    bullets(doc, [
        "主分支 `main` 保存稳定交付；功能分支命名 `feature/watchlist`、`feature/alert`。",
        "每次提交信息包含任务编号，例如 `docs: add SRS use cases`。",
        "文档、原型和代码都进入仓库，便于追踪评审修改。",
        "评审前打 tag：`v0.1-requirements`、`v0.2-prototype`、`v1.0-final`。"
    ])
    return save(doc, "StockMonitor-迭代计划.docx")


def build_evaluation():
    doc = Document()
    style_doc(doc, "迭代评估报告")
    add_history(doc)
    doc.add_heading("1. 评估摘要", level=1)
    doc.add_paragraph("本次迭代完成了项目定位、数据源调研、需求规约、用例规约、架构设计、HTML 界面原型和 GitHub 交付目录整理。系统目标由“股票盯盘软件”收敛为“Windows 端轻量自选股盯盘与提醒工具”。")
    doc.add_heading("2. 完成情况", level=1)
    add_table(doc, ["计划项", "状态", "说明"], [
        ["调研数据源可获得性", "完成", "明确 AKShare、BaoStock、公开行情接口与授权行情源边界。"],
        ["前景文档", "完成", "说明用户、问题、定位、竞争和约束。"],
        ["SRS", "完成", "定义功能需求、非功能需求、接口和原型映射。"],
        ["用例规约", "完成", "完成 6 个核心用例的基本流、备选流、前后置条件。"],
        ["软件构架文档", "完成", "完成分层、数据视图、部署视图和质量属性。"],
        ["界面原型", "完成", "HTML 原型可交互，覆盖主界面、行情、图表和提醒。"],
    ])
    doc.add_heading("3. 小组评审记录", level=1)
    add_table(doc, ["评审问题", "处理结果"], [
        ["原始需求包含 K 线但用户只需要趋势", "将 K 线简化为日收盘价折线，降低实现成本。"],
        ["公开行情接口稳定性不确定", "在 SRS 和架构中增加 Provider、缓存和 mock 数据方案。"],
        ["提醒功能可能连续弹窗", "规则增加触发状态，避免重复打扰。"],
        ["界面与用例映射不明显", "原型和 SRS 增加界面元素到需求编号的映射。"],
    ])
    doc.add_heading("4. 问题与改进", level=1)
    bullets(doc, [
        "PDF 作业要求存在加密或解析限制，后续提交前建议人工再核对课程细则。",
        "公开行情源适合课程演示，但商业发布必须确认数据授权。",
        "下一轮实现应优先补单元测试：提醒规则判断、行情适配器异常、SQLite 持久化。"
    ])
    doc.add_heading("5. 经验总结", level=1)
    bullets(doc, [
        "先做数据源可行性分析，可以避免需求写得过满而无法实现。",
        "用 Provider 接口隔离行情源，使系统在公开接口变化时仍能演示。",
        "商科用户更关心资产观察和风险提醒，图表应简洁清晰，而不是复刻专业交易终端。"
    ])
    doc.add_heading("6. 后续计划", level=1)
    numbered(doc, [
        "根据课堂反馈调整文档措辞和用例粒度。",
        "实现 Windows 客户端最小可运行版本。",
        "接入一个真实行情源并保留模拟数据降级。",
        "补充测试报告和最终演示脚本。"
    ])
    return save(doc, "StockMonitor-迭代评估报告.docx")


def main():
    paths = [
        build_vision(),
        build_srs(),
        build_use_cases(),
        build_architecture(),
        build_iteration_plan(),
        build_evaluation(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
